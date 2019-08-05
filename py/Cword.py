# -*- coding: utf-8 -*-
# @Time    : 2018/4/16 10:56
# @Author  : CRS
from docx import Document

path="C:/Users/Administrator/Desktop/登革热预警系统功能介绍.docx"
file = Document(path)
deal = []
for i in range(len(file.paragraphs)):
    if file.paragraphs[i].text != "":
        deal.append(file.paragraphs[i].text)

title = deal[0]
print(title)
last = deal[-1]
print(last)
deal.remove(deal[0])
deal.remove(deal[-1])
content = "".join(deal)
print(content)








