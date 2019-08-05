import codecs
import datetime
import operator
import os
import re
import jieba
import win32com.client
import time
from sklearn.feature_extraction.text import TfidfTransformer
from sklearn.feature_extraction.text import CountVectorizer
import matplotlib.pyplot as plt
import pymysql
from docx import Document

wdFormatDocument = 0
wdFormatDocument97 = 0
wdFormatDocumentDefault = 16
wdFormatDOSText = 4
wdFormatDOSTextLineBreaks = 5
wdFormatEncodedText = 7
wdFormatFilteredHTML = 10
wdFormatFlatXML = 19
wdFormatFlatXMLMacroEnabled = 20
wdFormatFlatXMLTemplate = 21
wdFormatFlatXMLTemplateMacroEnabled = 22
wdFormatHTML = 8
wdFormatPDF = 17
wdFormatRTF = 6
wdFormatTemplate = 1
wdFormatTemplate97 = 1
wdFormatText = 2
wdFormatTextLineBreaks = 3
wdFormatUnicodeText = 7
wdFormatWebArchive = 9
wdFormatXML = 11
wdFormatXMLDocument = 12
wdFormatXMLDocumentMacroEnabled = 13
wdFormatXMLTemplate = 14
wdFormatXMLTemplateMacroEnabled = 15
wdFormatXPS = 18
filelist = []

def JDBC():
    # 连接数据库
    conn = pymysql.connect(host="200.100.100.68", port=3306,
                               user="root", passwd="root", db="test", charset="utf8")
    return conn



# 对公文文件进行遍历
def getAllfileAndDirPath(sourcePath,*findType):
    """

    :param sourcePath: 需要遍历的文件路径
    :param findType: 需要找到的类型
    :return:
    """
    if not os.path.exists(sourcePath):
        return
    listName = os.listdir(sourcePath)
    for name in listName:
        absPath = os.path.join(sourcePath, name)
        if os.path.isfile(absPath):
            for type in findType:
                # 对文件进行过滤 absPath.find(".txt") != -1
                if absPath.endswith(type):
                    absPath1 = absPath.replace('\\', '/')
                    filelist.append(absPath1)
                # print('filePath:%s' % absPath)
        if os.path.isdir(absPath):
            getAllfileAndDirPath(absPath)

# 对文档进行分词处理
def fenci(file, ok):
    """

    :param file: 需要分词的文件
    :param ok: 分词处理后的文件
    :return:
    """
    # 正则匹配 找到文件名
    re_match = re.match(".*/(.*)?", file)
    if re_match:
        filname = re_match.group(1)
    pos = filname.find(".")
    # 在文件名后面加上时间戳
    filnames = filname[:pos] + str(time.time()).replace(".", "") + filname[pos:]
    # 读取文档
    try:
        f = codecs.open(file, 'r', encoding='GBK')
        last = codecs.open(os.path.join(ok, filnames), 'w', encoding='utf-8')
    except IOError:
        print("io error can't open !")
    # 加载停词库
    stopwords = stopwordslist()
    # 加载自定义词典
    # jieba.load_userdict('C:\\Users\\lenovo\\Desktop\\自定义词库.txt')
    # 对文档进行分词处理，采用精确模式
    value = re.compile(r'^[-+]?[0-9]+\.[0-9]+$')

    for line in f.readlines():
        seg_list = jieba.cut(line, cut_all=False)
        # 对空格，换行符进行处理
        for seg in seg_list:
            seg.encode("utf-8")
            if seg not in stopwords:
                #  需要修改
                if seg == '\r' or seg == "\n" or seg == "\r\n" or value.match(seg) :
                    #print(seg)
                    #print("打开分词存放文件")
                    pass
                else:
                    last.write(seg + " ")
    last.close()
    f.close()
    print("此文件分词结束")


# 进行TF-IDF计算
def Tfidf(okfile,sFilePath):
    """

    :param okfile: 需要计算的分词文件
    :param sFilePath: 处理完成后的文件路径
    :return:
    """
    corpus = []  # 存取文档的分词结果
    names = []
    for name in os.listdir(okfile):
        file = (os.path.join(okfile, name))
        pos = name.find(".")
        names.append(name[:pos])
        f = codecs.open(file, 'r+', encoding='utf-8')
        corpus.append(" ".join(f.readlines()))
        f.close()
    # 该类会将文本中的词语转换为词频矩阵，矩阵元素a[i][j] 表示j词在i类文本下的词频
    print(corpus)
    vectorizer = CountVectorizer()
    # 计算词语出现的次数
    # 将文本中的词语转换为词频矩阵
    x = vectorizer.fit_transform(corpus)
    print(x.toarray())
    # 该类会统计每个词语的tf-idf权值
    # max_df：这个给定特征可以应用在
    # tf - idf矩阵中，用以描述单词在文档中的最高出现率。假设一个词（term）在
    # 80 % 的文档中都出现过了，那它也许（在剧情简介的语境里）只携带非常少信息。
    # min_df：可以是一个整数（例如5）。意味着单词必须在
    # 5个以上的文档中出现才会被纳入考虑。在这里我设置为
    # 0.2；即单词至少在20 % 的文档中出现 。因为我发现如果我设置更小的
    # min_df，最终会得到基于姓名的聚类（clustering）——举个例子，好几部电影的简介剧情中老出现“Michael”或者“Tom”这些名字，然而它们却不携带什么真实意义。
    transformer = TfidfTransformer()
    # transformer.fit_transform是计算tf-idf
    tfidf = transformer.fit_transform(x)
    # 获取词袋中所有文本关键词
    word = vectorizer.get_feature_names()
    weight = tfidf.toarray()  # 对应的tfidf矩阵
    print("*" * 20 + "处理完所有文本")
    if not os.path.exists(sFilePath):
        os.mkdir(sFilePath)
    print(weight)
    # 这里将每份文档词语的TF-IDF写入tfidffile文件夹中保存
    for i in range(len(weight)):

        print("--------Writing all the tf-idf in the", i, u" file into ", names[i] +
              str(time.time()).replace(".", "") + '.txt', "--------")
        f = codecs.open(sFilePath + '/' + names[i] + str(time.time()).replace(".", "") + '.txt', 'w+', encoding='utf-8')
        # 排序
        sort = {}
        sort.clear()
        for j in range(len(word)):
            if weight[i][j] != 0.0:
                sort[word[j]] = str(weight[i][j])
        sorted_x = sorted(sort.items(), key=operator.itemgetter(1), reverse=True)
        for k, v in sorted_x:
            f.write(k + "  " + v + "\n")
        f.close()
    print("关键词提取结束！")
    #return weight


# 对文档转换格式
def change_file(filelist,typeNum,okpath,fileType=".txt"):
    """
    :param filelist: 需要转换格式的文档
    :param typeNum: 格式编号
    :param okpath: 另存为的文件路径
    :param fileType: 另存为的格式
    :return: 处理成功后的文件名
    """
    wordopen = win32com.client.Dispatch('Word.Application')
    # 后台运行，不显示，不警告
    wordopen.Visible = 0
    wordopen.DisplayAlerts = 0
    if os.path.exists(filelist):
        re_match = re.match(".*(/.*).*?", filelist)
        if re_match:
            filename = re_match.group(1)
        try:
            doc = wordopen.Documents.Open(filelist)
        except(Exception):
            print("error")
        lastname = okpath + filename.replace(".docx", fileType)
        pos = lastname.find(".")
        # 在文件名后面加上时间戳
        lastname = lastname[:pos] + str(time.time()).replace(".", "") + lastname[pos:]
        #print(lastname)
        try:
            # 乱码问题  待解决！！！！
            doc.SaveAs(lastname, typeNum)
        except(Exception):
            print("error")
        wordopen.Quit()
        return lastname
        # doc.close()

# 对文档进行处理 存数据库
def deal(path):
    """

    :param path: 需要处理的文件路径
    :return:
    """
    file = Document(path)
    re_match = re.match(".*/(.*)?", path)
    if re_match:
        filname = re_match.group(1)
    deal = []
    for i in range(len(file.paragraphs)):
        if file.paragraphs[i].text != "":
            deal.append(file.paragraphs[i].text)
    title = deal[0]
    #print(title)
    last = deal[-1]
    #print(last)
    deal.remove(deal[0])
    deal.remove(deal[-1])
    content = "".join(deal)
    dt = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    sql= "insert into document(title,content,end,time,name) VALUES (%s,%s,%s,%s,%s)"
    #sql1 = "select * from student"
    conn = JDBC()
    try:
        conn.cursor().execute(sql,(title,content,last,dt,filname,))
        conn.commit()
    except :
        print("失败！")
        conn.cursor().rollback()

# 停词库
def stopwordslist():
    """
    :return: 停词库
    """
    # 停词库
    stoppath = "C:/Users/Administrator/Desktop/chinese_stopword.txt"
    stopwordslist = [line.strip() for line in open(stoppath, 'r', encoding='utf-8').readlines()]
    # for _ in stopwordslist:
    #     print(_)
    return stopwordslist


def k_means(weight):
    'Start Kmeans:'
    from sklearn.cluster import KMeans
    clf = KMeans(n_clusters=4)  # 景区 动物 人物 国家
    s = clf.fit(weight)
    ''''' 
    print 'Start MiniBatchKmeans:' 
    from sklearn.cluster import MiniBatchKMeans 
    clf = MiniBatchKMeans(n_clusters=20) 
    s = clf.fit(weight) 
    print s 
    '''

    # 中心点
    print(clf.cluster_centers_)

    # 每个样本所属的簇
    label = []  # 存储1000个类标 4个类
    print(clf.labels_)
    i = 1
    while i <= len(clf.labels_):
        label.append(clf.labels_[i - 1])
        i = i + 1

        # 用来评估簇的个数是否合适，距离越小说明簇分的越好，选取临界点的簇个数  958.137281791
    print(clf.inertia_)

    ########################################################################
    #                               第三步 图形输出 降维

    from sklearn.decomposition import PCA
    pca = PCA(n_components=2)  # 输出两维
    newData = pca.fit_transform(weight)  # 载入N维
    # 5A景区
    x1 = []
    y1 = []
    i = 0
    while i < 13:
        x1.append(newData[i][0])
        y1.append(newData[i][1])
        i += 1

        # 动物
    x2 = []
    y2 = []
    i = 400
    while i < 23:
        x2.append(newData[i][0])
        y2.append(newData[i][1])
        i += 1

        # 人物
    x3 = []
    y3 = []
    i = 600
    while i < 33:
        x3.append(newData[i][0])
        y3.append(newData[i][1])
        i += 1

        # 国家
    x4 = []
    y4 = []
    i = 800
    while i < 55:
        x4.append(newData[i][0])
        y4.append(newData[i][1])
        i += 1

        # 四种颜色 红 绿 蓝 黑
    plt.plot(x1, y1, 'or')
    plt.plot(x2, y2, 'og')
    plt.plot(x3, y3, 'ob')
    plt.plot(x4, y4, 'ok')
    plt.show()


if __name__ == "__main__":

    getAllfileAndDirPath("C:/Users/Administrator/Desktop/智源文档/公函分析/数据源文件",".txt")
ok = "C:/Users/Administrator/Desktop/智源文档/公函分析/分词"
sFilePath = 'C:/Users/Administrator/Desktop/智源文档/公函分析/关键词提取结果'
print("进入循环")
for file in filelist:
    print(file)
    #print(change_file(file,wdFormatText,'C:/Users/Administrator/Desktop/智源文档/公函分析/数据源文件'))
    # print("*"*20+"分词开始")
    fenci(file,ok)
    # print("开始储存")
    # deal(file)

Tfidf(ok,sFilePath)
#conn.close()


# fenci(file, ok)
# print("分词结束")
# print("-" * 20 + "\n" + "提取关键词开始")
# x = Tfidf(ok)
# print("聚类分析")
# k_means(x)
