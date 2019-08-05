# -*- coding: utf-8 -*-
# @Time    : 2018/11/16 16:55
# @Author  : Crsboom
# @Email   : 1105959249@qq.com
# @File    : saveWord.py
# @Software: PyCharm
# %%       最终报告结果输出

import os
from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor
import cx_Oracle
from util.jdbc import Oracle_conn,Oracle_cur

def Generate_word():
    results = {'1':1,'2':2,'3':3,'4':4,'5':5,'6':6,'7':7,'8':8,'9':9}
    # 1打开文档
    document = Document()
    # 2加入文字
    document.add_heading('直接测定法结果展示', 0)
    heading1 = document.add_heading()
    heading1 = heading1.add_run(u'直接测定法')
    p = document.add_paragraph()
    run = p.add_run(u'''    直接测定法是直接测得药物对各个动物最小销量或最小致死量的鉴定方法。\n\
                     XT和XS为T和S组各只动物的对数最小致死量；它们的均值XS和XT为S和T的等反应剂量；
                     NS和NT为S和T组的动物数。''')
    heading2 = document.add_heading()
    heading2 = heading2.add_run(u'符号解释')
    p1 = document.add_paragraph()
    run1 = p1.add_run(u'''S和T :生物实验体\n\
                        R：等反应剂量比\n\
                        DS,DT：等反应剂量\n\
                        M：S和T的对数等反应剂量\n\
                        PT：T的测得效价\n\
                        AT：估计效价\n\
                        S2：指从实验结果的变异种分去不同剂量及不同因级对变异的影响后，剩余的变异成分\n\
                        FL：可信限，标志鉴定结果和精密度\n\
                        SM：M的标准误''')
    heading3 = document.add_heading()
    heading3 = heading3.add_run(u'结果展示')
    table = document.add_table(rows=len(results), cols=2)
    results_keys = list(results.keys())
    for i in range(len(results)):
        hdr_cells = table.rows[i].cells
        hdr_cells[0].text = results_keys[i]
        hdr_cells[1].text = str(results[results_keys[i]])
    #if results['R的FL% (%)'] > 15 or results['PT的FL% (%)'] > 15:
    p2 = document.add_paragraph()
    run2 = p2.add_run(u'注：本法的可信限FL%不得大于15%')
    run2.italic = True
    run2.bold = True
    run2.font.color.rgb = RGBColor(0xFF, 0x66, 0x00)
    for i in [heading1, run, heading2, run1, heading3]:
        i.font.name = u'宋体'
        r = i._element
        r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        ##word中生成的表格格式还不太会弄，不知道为啥同一句代码，别人生成的表格是有线的，我却没
    # 3保存文件
    algorithm = ['直接测定法']
    # document.save(algorithm[a]+'结果'+results_fname+'.docx')
    document.save('结果.docx')

def ins_Oracle(file):
    f = open(file, 'rb')
    conn = cx_Oracle.connect('detail','detail','200.100.100.68:1521/CRS')
    cur = conn.cursor()
    sql = '''insert into ib_tbs_detailedinf(fdiseq,fbiseq,forgno,fspno,fdatno,finfseq,ffilenm,fcontent,fpath,fopdt,fempid,fhiino)
            values(88888888,0,1,'1','01',1,'我是导入的.docx',:blobData,NULL,sysdate,'dsj','1')'''
    print(sql)
    cur.setinputsizes(blobData=cx_Oracle.BLOB)
    try:
        cur.execute(sql,{'blobData':f.read()})
        cur.execute('commit')
        f.close()
    except Exception as e:
        print(e)
    cur.close()
    conn.close()

def get_depend(id):
    sql = '''select c1,c2,c3,c4,c5,c6 from lab_form where fid = 'aa61a8dc-ebd4-11e8-af48-f0761cbe1239' '''



if __name__ == '__main__':
    # Generate_word()
    ins_Oracle('结果.docx')