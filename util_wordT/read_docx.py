# -*- coding: utf-8 -*-
# @Time    : 2019/2/20 13:23
# @Author  : Crsboom
# @Email   : 1105959249@qq.com
# @File    : read_docx.py
# @Software: PyCharm
from util.googole_translate import translate
from docx import Document
from docx.oxml.ns import qn



path=r"G:\工作总结及进度\缓冲目录\翻译书\E-I.docx"
file = Document(path)
deal = []
document = Document()
document.styles['Normal'].font.name = u'宋体'
document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

for i in range(len(file.paragraphs)):
    if file.paragraphs[i].text != "":
        #deal.append(file.paragraphs[i].text)
        text = file.paragraphs[i].text.split('.')
        print(file.paragraphs[i].text)
        document.add_paragraph().add_run(file.paragraphs[i].text)
        x = []
        for i in text:
            # print(i)
            res = translate.translate(i)
            x.append(res)
        print("".join(x))
        document.add_paragraph().add_run("".join(x))

document.save('result.docx')




# title = deal[0]
# print(title)
# last = deal[-1]
# print(last)
# deal.remove(deal[0])
# deal.remove(deal[-1])
# content = "".join(deal)
# print(content)
