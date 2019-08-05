# -*- coding: utf-8 -*-
# @Time    : 2019/3/1 16:32
# @Author  : Crsboom
# @Email   : 1105959249@qq.com
# @File    : pdftotxt.py
# @Software: PyCharm
from urllib.request import urlopen
#from urllib import urlopen
from io import StringIO
from pdfminer.pdfinterp import PDFResourceManager, process_pdf
from pdfminer.converter import TextConverter
from pdfminer.layout import  LAParams
import os
from docx import Document


# 读取pdf的函数，返回内容
def readPdf(path,pdf_file):
    f = open(os.path.join(path,pdf_file), 'rb')
    rsrcmgr = PDFResourceManager()
    retstr = StringIO()
    laparams = LAParams()
    device = TextConverter(rsrcmgr=rsrcmgr, outfp=retstr, laparams=laparams)
    process_pdf(rsrcmgr=rsrcmgr, device=device, fp=f)
    device.close()
    content = retstr.getvalue()
    retstr.close()
    saveTXT(pdf_file,content)


def getPDF(path):
    file_list = os.listdir(path)
    num = 1
    for i in file_list:
        if i.find('.pdf') != -1:
            print('第 %s 个pdf文件：%s'%(num,i))
            num += 1
            readPdf(path,i)


def saveTXT(file_name,content):
    pos = file_name.find('.')
    res = file_name[0:pos]+'.txt'
    sFilePath = './txt'
    if not os.path.exists(sFilePath):
        os.mkdir(sFilePath)
    with open( os.path.join(sFilePath,res) , 'w', encoding='utf8') as f:
        f.write(content)
        print('写入成功----文件名为：%s'% res)


def word_txt(path):
    file = Document(path)
    deal = []
    for i in range(len(file.paragraphs)):
        if file.paragraphs[i].text != "":
            deal.append(file.paragraphs[i].text)
    print(deal)


if __name__ == '__main__':
    # content = readPdf(f)
    # print(type(content))
    # with open('123.txt','w',encoding='utf8') as f:
    #     f.write(content)
    #print(content)
    # getPDF(r'G:\工作总结及进度\招标数据\zhaobiaofile')
    word_txt('result.docx')


